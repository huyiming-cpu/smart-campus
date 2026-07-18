"""生成 3.项目开发计划.docx — 智慧校园服务平台"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ============ 封面 ============
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('智慧校园服务平台')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('项目开发计划')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'

for _ in range(2):
    doc.add_paragraph()

add_table(
    ['项目编号: SWJTU-SC-2026-001\n智慧校园服务平台', '项目开发计划\nVersion: 1.0\n\n项目承担部门：软件工程实习项目组\n撰写人（签名）：__________\n完成日期：2026-07-17\n本文档使用部门：■主管领导 ■项目组\n□客户（市场） □维护人员 □用户\n评审负责人（签名）：__________\n评审日期：__________', ''],
    []
)

doc.add_page_break()

# ============ 修订历史 ============
add_heading('修订文档历史记录', 1)
add_table(
    ['日期', '版本', '说明', '作者'],
    [
        ['2026-07-10', '0.1', '草稿，完成项目初期规划', '项目组'],
        ['2026-07-15', '1.0', '正式发布，包含完整功能计划', '项目组'],
        ['2026-07-17', '1.1', '补充智慧校园全部功能模块详细计划', '项目组'],
    ]
)
doc.add_paragraph()

# ============ 1. 前言 ============
add_heading('1. 前言', 1)
add_heading('1.1 目的', 2)
add_para('本计划应用于项目《智慧校园服务平台》开发的整个生命周期。旨在明确项目目标、范围、角色职责、开发过程、工作任务分解及各项管理计划，为项目组全体成员提供统一的行动指南，确保项目按期、高质量交付。')

add_heading('1.2 术语与缩略语', 2)
add_para('JWT：JSON Web Token，用于用户身份认证的无状态令牌')
add_para('RBAC：Role-Based Access Control，基于角色的访问控制模型')
add_para('SPA：Single Page Application，单页面应用程序')
add_para('BCrypt：密码哈希加密算法')
add_para('CRUD：Create / Read / Update / Delete，数据基本操作')
add_para('ORM：Object-Relational Mapping，对象关系映射')
add_para('PPQA：Process and Product Quality Assurance，过程和产品质量保证')
add_para('CM：Configuration Management，配置管理')

# ============ 2. 项目概述 ============
add_heading('2. 项目概述', 1)

add_heading('2.1 项目背景和目标', 2)
add_para('随着高校信息化建设的深入推进，传统校园管理方式面临信息孤岛、流程繁琐、师生体验差等痛点。学院需要一个统一的信息化平台，整合教务管理、一卡通、缴费、评教、科研、公文等核心业务，同时为管理员提供高效的管理工具，为领导层提供数据决策支持。')
add_para('项目目标：', bold=True)
add_para('1. 建成覆盖学生、教师、领导、管理员四类角色的智慧校园服务平台；')
add_para('2. 实现10大功能模块共60项细分功能，涵盖教学、财务、行政、科研等全业务场景；')
add_para('3. 集成AI智能助手，为全校用户提供"问答即服务"的智能体验；')
add_para('4. 系统支持日均1000+并发用户，核心接口响应时间≤500ms；')
add_para('5. 采用前后端分离架构，满足可维护性、可扩展性要求。')

add_heading('2.2 项目范围', 2)
add_para('本项目覆盖智慧校园的全部核心业务场景：')
add_para('• 用户权限：登录注册、个人信息管理、角色权限管理(RBAC)、注册审核')
add_para('• 通知公告：系统公告发布与管理、学校新闻发布与管理')
add_para('• 课程成绩：课程管理、课表安排、学生选课、成绩录入与查询、补考重修管理')
add_para('• 考试评教：期末考试安排、等级考试报名、教学评价')
add_para('• 一卡通：校园卡信息查询、消费记录管理、学费缴费管理')
add_para('• 学生服务：奖学金/助学金/贷款申请、竞赛报名、心理咨询预约、实践项目预约')
add_para('• 教师发展：科研项目管理、毕业设计管理')
add_para('• 公文计划：公文审批流转、工作计划管理')
add_para('• 系统管理：仪表盘、系统配置、日志审计、固定资产、实验室管理')
add_para('• AI助手：功能导航引导、个人数据查询、业务指引、统计报表查询')
add_para('项目不包含：在线支付功能、第三方地图集成、视频会议功能。')

add_heading('2.3 交付的产品', 2)
add_table(
    ['所属阶段', '交付工件名称', '工件类型', '预定交付日期'],
    [
        ['项目策划', '项目开发计划', 'Word文档', '2026-07-10'],
        ['需求分析', '软件需求规约', 'Word文档', '2026-07-12'],
        ['系统设计', '系统架构设计说明书', 'Word文档', '2026-07-17'],
        ['系统设计', '数据库设计说明书', 'Word文档', '2026-07-17'],
        ['编码实现', '智慧校园服务平台源码(后端+前端)', 'Java/Vue源码', '2026-07-28'],
        ['测试', '测试计划', 'Word文档', '2026-07-25'],
        ['测试', '测试用例', 'Word文档', '2026-07-27'],
        ['测试', '测试日志', 'Word文档', '2026-07-30'],
        ['交付', '项目开发总结报告', 'Word文档', '2026-08-01'],
        ['交付', '用户操作手册', 'Word文档', '2026-08-01'],
        ['交付', '答辩PPT', 'PPTX', '2026-08-05'],
    ]
)
doc.add_paragraph()

add_heading('2.4 约束和假设', 2)
add_para('约束条件：')
add_para('• 项目须于2026年8月5日前完成全部开发和文档交付')
add_para('• 数据库使用MySQL 8.4.5，后端使用Spring Boot 4.1.0 + JDK 17')
add_para('• 前端使用Vue 3.5 + Vite 8 + Element Plus')
add_para('• 开发周期约4周，项目组成员需全程参与')
add_para('假设条件：')
add_para('• 核心成员不被调离项目组')
add_para('• 开发环境和服务器资源按时到位')
add_para('• 需求不会发生重大变更')

# ============ 3. 角色和职责 ============
add_heading('3. 角色和职责', 1)
add_heading('3.1 利益相关人角色和职责', 2)
add_table(
    ['姓名', '角色', '在项目中作用'],
    [
        ['[姓名1]', '项目经理', '项目管理、进度跟踪、风险控制、与导师沟通'],
        ['[姓名2]', '系统架构师/后端开发', '系统架构设计、后端核心模块开发、数据库设计'],
        ['[姓名3]', '后端开发', '后端业务模块开发、单元测试、API文档编写'],
        ['[姓名4]', '前端开发', '前端架构设计、页面开发、前后端联调'],
        ['[姓名5]', '前端/测试工程师', '前端页面开发、测试用例编写、功能测试执行'],
        ['[姓名6]', '文档/测试工程师', '各类文档撰写、测试执行、Bug跟踪'],
    ]
)
doc.add_paragraph()

add_heading('3.2 利益相关人介入计划', 2)
add_table(
    ['相关部门', '相关角色', '相关人员', '备注'],
    [
        ['指导老师', '项目导师', '[导师姓名]', '需求确认、进度检查、技术指导'],
        ['项目组', 'PM', '[项目经理]', '资源协调与保障'],
        ['质量保证', 'PPQA检查员', '[QA人员]', '过程质量检查'],
        ['测试组', '测试负责人', '[测试人员]', '测试计划与执行'],
    ]
)
doc.add_paragraph()

# ============ 4. 项目已定义过程 ============
add_heading('4. 项目的已定义过程', 1)
add_heading('4.1 项目生命周期选择', 2)
add_para('本项目采用迭代式增量开发模型，参考RUP统一过程框架，划分为先启、精化、构建、交付四个阶段。各阶段迭代进行，每个迭代产出可运行的软件增量。')

add_heading('4.2 项目阶段划分及主要工作产品', 2)
add_table(
    ['所属阶段', '交付工件名称', '工件类型', '预定交付日期'],
    [
        ['先启阶段（第1周）', '项目开发计划', 'Word文档', '2026-07-10'],
        ['先启阶段（第1周）', '软件需求规约', 'Word文档', '2026-07-12'],
        ['精化阶段（第2周）', '系统架构设计说明书', 'Word文档', '2026-07-17'],
        ['精化阶段（第2周）', '数据库设计说明书', 'Word文档', '2026-07-17'],
        ['构建阶段（第3-4周）', '全部功能模块源码', '源码+可运行程序', '2026-07-28'],
        ['构建阶段（第3-4周）', '测试计划/测试用例', 'Word文档', '2026-07-27'],
        ['交付阶段（第4周）', '测试日志/总结报告/答辩PPT', 'Word/PPTX', '2026-08-05'],
    ]
)
doc.add_paragraph()

add_heading('4.3 本项目采用的过程', 2)
add_table(
    ['工程 Engineering', '度量与分析 MA', '过程管理 ProcessManage', '项目管理 ProjectManage', '支持 Support'],
    [
        ['需求管理(REQM)\n需求开发(RD)', '度量与分析(MA)', '组织过程定义(OPD)', '项目策划(PP)', '过程和产品质量保证(PPQA)'],
        ['技术解决方案(TS)', '', '组织过程焦点(OPF)', '项目监督和控制(PMC)', '配置管理(CM)'],
        ['产品集成(PI)', '', '组织级培训(OT)', '风险管理(RSKM)', '决策分析和解决方案(DAR)'],
        ['', '', '', '集成项目管理(IPM)', ''],
    ]
)
doc.add_paragraph()

add_heading('4.4 裁剪结论', 2)
add_para('本项目采用生命周期阶段裁剪方式。针对4周的开发周期，适当简化了文档模板中的部分过程域要求，保留核心的工程类和管理类过程，裁剪了部分适用于大型项目的支持类过程。')

# ============ 5. 工作任务分解 ============
add_heading('5. 工作任务分解（WBS）', 1)
add_para('一级任务：项目策划 → 需求分析 → 系统设计 → 编码实现(10个模块) → 测试 → 文档与交付')
add_para('')
add_para('二级任务分解（编码阶段）：', bold=True)
add_table(
    ['任务编号', '任务名称', '负责人', '工作量(人天)', '前置任务'],
    [
        ['WBS-01', '项目环境搭建（后端+前端+数据库）', '架构师', '2', '-'],
        ['WBS-02', 'user模块-登录注册个人信息', '后端开发A', '2', 'WBS-01'],
        ['WBS-03', 'notice模块-公告新闻管理', '后端开发A', '1.5', 'WBS-02'],
        ['WBS-04', 'system模块-仪表盘配置资产', '后端开发A', '2', 'WBS-02'],
        ['WBS-05', 'document模块-公文计划管理', '后端开发B', '1.5', 'WBS-02'],
        ['WBS-06', 'card模块-一卡通缴费管理', '后端开发B', '2', 'WBS-02'],
        ['WBS-07', 'teacher模块-科研毕设管理', '后端开发B', '1.5', 'WBS-02'],
        ['WBS-08', 'student模块-学生服务管理', '后端开发A', '2', 'WBS-06'],
        ['WBS-09', 'course模块-课程成绩管理', '后端开发B', '2.5', 'WBS-06'],
        ['WBS-10', 'exam模块-考试评教管理', '后端开发A', '2', 'WBS-07'],
        ['WBS-11', 'ai模块-智能助手', '后端开发B', '2', 'WBS-09'],
        ['WBS-12', 'user扩展-用户管理RBAC统计', '后端开发A', '2.5', 'WBS-10'],
        ['WBS-13', '前端-页面组件开发(4角色×10模块)', '前端开发A/B', '5', 'WBS-03'],
        ['WBS-14', '前后端联调', '全员', '2', 'WBS-11, WBS-13'],
        ['WBS-15', '单元测试+集成测试', '测试工程师', '3', 'WBS-14'],
        ['WBS-16', '文档撰写与PPT制作', '文档工程师', '3', 'WBS-14'],
    ]
)

# ============ 6. 项目估计 ============
add_heading('6. 项目估计', 1)
add_para('• 总工作量估计：约 35 人天（6人×4周）')
add_para('• 代码量估计：后端约8000行Java代码 + 前端约5000行Vue/JS代码')
add_para('• 数据库表：37张')
add_para('• API接口数：约80个RESTful端点')
add_para('• 前端页面数：约35个Vue页面组件')

# ============ 7. 技能和培训 ============
add_heading('7. 项目所需技能和培训计划', 1)
add_heading('7.1 项目所需技能', 2)
add_table(
    ['角色', '预计人数', '到位时间', '技能/经验要求'],
    [
        ['项目经理', '1', '2026-07-08', '2年以上项目管理经验，熟悉敏捷开发流程'],
        ['系统架构师', '1', '2026-07-08', '3年以上Java Web开发经验，熟悉Spring Boot、Vue'],
        ['后端开发工程师', '2', '2026-07-08', '1年以上Spring Boot/JPA开发经验'],
        ['前端开发工程师', '2', '2026-07-08', '1年以上Vue3 + Element Plus开发经验'],
        ['测试工程师', '1', '2026-07-15', '了解Web应用测试方法和工具'],
    ]
)
doc.add_paragraph()

add_heading('7.2 培训计划', 2)
add_table(
    ['培训时间', '培训内容', '培训方式', '参加人员'],
    [
        ['2026-07-08', '项目需求讲解与功能演示', '会议讲解', '全员'],
        ['2026-07-09', 'Spring Boot 4.1新特性介绍', '文档+演示', '后端组'],
        ['2026-07-09', 'Vue 3.5 + Vite 8开发入门', '文档+演示', '前端组'],
        ['2026-07-10', 'Git版本控制与协作规范', '实操演练', '全员'],
    ]
)

# ============ 8. 监控计划 ============
add_heading('8. 项目监控计划', 1)
add_heading('8.1 活动安排', 2)
add_para('• 每日站会（15分钟）：每人汇报昨天完成、今天计划、遇到的阻碍')
add_para('• 每周周会（30分钟）：回顾本周进度、评估风险、调整下周计划')
add_para('• 每周由项目经理填写《项目进度跟踪表》和《项目风险跟踪表》')
add_para('• 项目完成后，由项目经理完成《项目开发总结报告》')

add_heading('8.2 偏差控制', 2)
add_table(
    ['监控参数', '控制值', '行动'],
    [
        ['工作量偏差', '阈值20%，预警值15%', '分析原因，调整人员或范围'],
        ['进度偏差', '阈值15%，预警值10%', '评估关键路径，决定是否赶工'],
        ['Bug数量', '严重Bug>5个', '暂停新功能开发，集中修复'],
    ]
)

# ============ 9. 工作环境 ============
add_heading('9. 工作环境', 1)
add_heading('9.1 开发环境', 2)
add_table(
    ['类别', '工具/技术', '版本', '用途'],
    [
        ['IDE', 'IntelliJ IDEA', '2026.1', '后端Java开发'],
        ['IDE', 'VS Code', 'latest', '前端Vue开发'],
        ['JDK', 'JDK 17', '17.0.x', 'Java运行环境'],
        ['构建工具', 'Maven Wrapper', '3.9.x', '后端项目构建管理'],
        ['构建工具', 'Vite', '8.x', '前端构建工具'],
        ['数据库', 'MySQL', '8.4.5', '关系型数据库'],
        ['数据库工具', 'DBeaver', '24.x', '数据库管理GUI'],
        ['版本控制', 'Git', 'latest', '代码版本管理'],
        ['API测试', 'Postman / curl', 'latest', '接口测试'],
        ['文档', 'Microsoft Word', '2024+', '文档撰写'],
    ]
)
doc.add_paragraph()

add_heading('9.2 测试环境', 2)
add_table(
    ['类别', '配置', '说明'],
    [
        ['后端服务器', 'Windows 11 + JDK 17', '本地开发机'],
        ['前端服务器', 'Vite Dev Server (端口5173)', '开发热更新'],
        ['数据库服务器', 'MySQL 8.4.5 (端口3306)', '本地数据库实例'],
        ['浏览器测试', 'Chrome 130+, Edge 130+, Firefox 120+', '兼容性测试'],
    ]
)

# ============ 保存 ============
output_path = r'D:\software\smart-campus\3.项目开发计划-智慧校园.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
