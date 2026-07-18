"""生成 5.系统架构设计说明书.docx — 智慧校园服务平台"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

# ============ 封面 ============
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('智慧校园服务平台')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('系统架构设计说明书')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'

for _ in range(2):
    doc.add_paragraph()

add_table(
    ['', '智慧校园服务平台\n系统架构设计说明书\nVersion: 1.0\n\n项目承担部门：软件工程实习项目组\n撰写人（签名）：__________\n完成日期：2026-07-17\n本文档使用部门：■主管领导 ■项目组\n□客户（市场） ■维护人员 □用户\n评审负责人（签名）：__________\n评审日期：__________', ''],
    []
)
doc.add_page_break()

# ============ 修订历史 ============
add_heading('修订文档历史记录', 1)
add_table(
    ['日期', '版本', '说明', '作者'],
    [
        ['2026-07-10', '0.1', '草稿，完成系统架构初稿', '项目组'],
        ['2026-07-15', '0.5', '补充核心模块详细设计', '项目组'],
        ['2026-07-17', '1.0', '正式发布，覆盖全部功能模块', '项目组'],
    ]
)
doc.add_paragraph()

# ============ 1. 系统概述 ============
add_heading('1. 系统概述', 1)

add_heading('1.1 项目背景与目标', 2)
add_para('随着高校信息化建设从"数字化"向"智慧化"迈进，传统校园管理系统存在信息孤岛严重、业务协同困难、用户体验不一致等痛点。智慧校园服务平台的定位是面向全校师生的一站式信息服务平台，整合教务管理、一卡通、缴费、评教、科研、公文、AI助手等10大核心业务模块，覆盖学生、教师、领导、管理员四类角色的60项功能需求。')
add_para('系统建设目标：提供统一入口的校园服务门户（PC Web端）；实现四角色权限管理（RBAC），数据隔离、功能隔离；集成AI智能助手，实现"问答即服务"的用户体验；采用前后端分离架构，保障系统可维护性和可扩展性。')

add_heading('1.2 系统定位', 2)
add_para('• 核心角色：学生（课程查询、成绩查看、一卡通、评教、申请服务）、教师（任课管理、成绩录入、科研管理、毕设指导）、领导（数据统计、公文审批、工作计划）、管理员（用户管理、系统配置、数据维护）')
add_para('• 覆盖范围：涵盖"教学管理—校园生活—行政办公—数据决策"全业务链，支撑"学、教、管、评"一体化。')

add_heading('1.3 核心需求摘要', 2)
add_para('功能需求：10大功能模块（用户权限、通知公告、课程成绩、考试评教、一卡通缴费、学生服务、教师发展、公文计划、系统管理、AI助手），共计60项细分功能，覆盖学生15项、教师14项、领导13项、管理员18项。', bold=False)
add_para('非功能需求：')
add_para('• 易用性：响应式布局，操作步骤≤3步，关键任务有引导提示')
add_para('• 性能：支持日均500+并发用户，核心接口响应时间≤500ms')
add_para('• 兼容性：支持Chrome 130+、Edge 130+、Firefox 120+')
add_para('• 安全性：JWT无状态认证、BCrypt密码加密、SQL参数化防注入、操作日志可追溯')

doc.add_page_break()

# ============ 2. 总体架构设计 ============
add_heading('2. 总体架构设计', 1)

add_heading('2.1 架构分层（纵向）', 2)
add_para('系统采用经典的四层B/S架构，前后端通过RESTful API通信：')

add_para('1. 前端层（Presentation Layer）：', bold=True)
add_para('• 技术栈：Vue 3.5 + Vite 8 + Element Plus + Axios + Pinia + Vue Router')
add_para('• 按角色组织路由：/student/*、/teacher/*、/leader/*、/admin/*')
add_para('• Pinia状态管理：存储用户Token、角色信息、全局配置')
add_para('• Axios拦截器：请求自动携带JWT Token，响应异常统一处理')

add_para('2. 接入层（Gateway Layer）：', bold=True)
add_para('• Vite开发代理：开发环境 /api → localhost:8080（Spring Boot后端）')
add_para('• 生产环境由Nginx反向代理，统一处理跨域和静态资源')

add_para('3. 应用层（Application Layer — Spring Boot）：', bold=True)
add_para('• 控制层(Controller)：接收HTTP请求，参数校验后调用Service')
add_para('• 业务层(Service)：核心业务逻辑、事务管理、权限校验')
add_para('• 持久层(Repository)：基于Spring Data JPA，封装数据库操作')
add_para('• 安全层(Security)：JwtFilter拦截所有非/auth请求，验证Token有效性')
add_para('• 10个业务模块按包组织，模块间通过Service互相调用，避免循环依赖')

add_para('4. 数据层（Data Layer）：', bold=True)
add_para('• MySQL 8.4.5：存储用户、课程、成绩、一卡通等37张业务表')
add_para('• HikariCP连接池：Spring Boot默认，最大连接数20')
add_para('• JPA自动建表：spring.jpa.hibernate.ddl-auto=update')

add_heading('2.2 核心模块划分（横向）', 2)
add_table(
    ['模块名称', '核心职责', '关键使用者', '涉及数据库表'],
    [
        ['user（用户权限）', '登录注册、个人信息、RBAC权限、用户管理、注册审核', '全部角色', 'user, user_profile, role, user_role, role_permission, permission, identity_registry, register_application, login_log, password_history (10张)'],
        ['notice（通知公告）', '公告CRUD与置顶、新闻CRUD与分类筛选、角色定向推送', '全部角色', 'announcement, news (2张)'],
        ['course（课程成绩）', '课程管理、课表安排、学生选课、成绩录入与查询、补考重修', '学生、教师、管理员', 'course, course_schedule, student_course, student_retake, score_change_log (5张)'],
        ['exam（考试评教）', '期末考试安排、等级考试(CET4/6等)报名、在线评教', '学生、教师、管理员', 'exam_arrangement, level_exam, level_exam_registration, teaching_evaluation (4张)'],
        ['card（一卡通）', '校园卡余额查询、消费流水、学费缴费与欠费管理', '学生、管理员', 'card, card_transaction, payment_record (3张)'],
        ['student（学生服务）', '奖/助/贷/困难生申请、竞赛报名、心理咨询预约、实践项目', '学生、教师(辅导员)', 'student_application, competition_registration, consultation_appointment, practice_project_registration (4张)'],
        ['teacher（教师发展）', '科研项目管理、论文统计、毕业设计选题与指导', '教师、领导', 'scientific_research, graduation_design (2张)'],
        ['document（公文计划）', '公文提交与审批、工作计划安排与跟踪', '教师、领导', 'document, work_plan (2张)'],
        ['system（系统管理）', '仪表盘、系统配置、日志审计、固定资产、实验室管理、领导统计报表', '领导、管理员', 'sys_config, fixed_asset, lab, login_log (4张)'],
        ['ai（AI助手）', '知识库管理、意图识别、多轮对话、角色自动感知', '全部角色', 'ai_knowledge_base, ai_conversation_history (2张)'],
    ]
)
doc.add_paragraph()

add_heading('2.3 系统边界', 2)
add_para('包含范围：以上10个模块的全部60项功能，以及JWT认证、Vite代理、全局异常处理等基础设施。')
add_para('排除范围：在线支付网关集成、第三方地图/导航服务、视频会议/直播功能、移动端App（本期仅Web端）。')

doc.add_page_break()

# ============ 3. 技术选型 ============
add_heading('3. 技术选型设计', 1)
add_table(
    ['技术类别', '选型方案', '选型理由', '版本'],
    [
        ['前端框架', 'Vue 3 + Vite + Element Plus', 'Vue 3 Composition API性能优、Element Plus组件丰富、Vite构建速度快', 'Vue 3.5 / Vite 8 / Element Plus 2.x'],
        ['前端状态管理', 'Pinia', 'Vue 3官方推荐，TypeScript友好，比Vuex更轻量', 'Pinia 2.x'],
        ['前端路由', 'Vue Router', 'Vue 3官方路由方案，支持角色路由守卫', 'Vue Router 4.x'],
        ['HTTP客户端', 'Axios', '支持拦截器、Promise API，方便封装', 'Axios 1.x'],
        ['后端框架', 'Spring Boot', 'Java生态最主流微服务框架，自动配置、starter丰富', 'Spring Boot 4.1.0'],
        ['ORM', 'Spring Data JPA / Hibernate', '自动建表、命名查询、减少SQL编写量', 'Hibernate 6.x'],
        ['安全认证', 'Spring Security + JWT (jjwt)', '无状态Token认证，适合前后端分离', 'jjwt 0.12.x'],
        ['密码加密', 'BCrypt', '自适应哈希算法，抗暴力破解', 'Spring Security内置'],
        ['数据库', 'MySQL', '成熟稳定、事务支持、社区活跃', 'MySQL 8.4.5'],
        ['构建工具', 'Maven Wrapper', '无需预装Maven，统一构建环境', 'Maven 3.9.x'],
        ['前端构建', 'Vite', 'ESBuild预构建，HMR热更新极快', 'Vite 8.x'],
    ]
)
doc.add_paragraph()

# ============ 4. 核心模块详细设计 ============
add_heading('4. 核心模块详细设计', 1)
add_heading('4.1 以"用户登录认证"模块为例', 2)

add_heading('4.1.1 核心业务流程', 2)
add_para('登录流程：', bold=True)
add_para('1. 前端：用户输入学号/工号 + 密码 → Axios POST /auth/login')
add_para('2. 后端：UserController接收请求 → UserService.login()')
add_para('3. 业务层：a) 根据identity_number查user表 → b) BCrypt比对密码 → c) 检查账号状态（NORMAL/LOCKED/DELETED）→ d) 查user_role获取角色 → e) 生成JWT Token（含userId、identityNumber、roleCode）')
add_para('4. 响应：返回Token + 用户基本信息（姓名、角色编码、角色名称）')
add_para('5. 前端：Pinia存储Token + 角色信息 → 路由守卫按角色跳转到对应首页')

add_para('注册流程：', bold=True)
add_para('1. 前端：用户填写学号+姓名+身份证后6位+密码+联系方式')
add_para('2. 后端：UserService.register()')
add_para('3. 身份核验：查identity_registry表，匹配identity_number + name + id_card_last6，校验status=0（未注册）')
add_para('4. 创建用户：user表新增记录 → BCrypt加密密码 → user_role表关联角色 → user_profile表创建档案（根据年级字段自动计算入学/毕业时间：如"2023级" → 2023-09-01 ~ 2027-06-01）')
add_para('5. 更新注册状态：identity_registry.status置为1，记录联系电话')

add_heading('4.1.2 核心数据表设计', 2)
add_table(
    ['表名', '核心字段', '字段说明'],
    [
        ['user', 'id, identity_number, password_hash, name, gender, phone, email, status, last_login_time', '用户主表，password_hash存BCrypt密文，status含NORMAL/LOCKED/DELETED'],
        ['user_role', 'id, user_id, role_id, is_default', '用户-角色关联表，is_default=1表示默认角色'],
        ['role', 'id, role_code, role_name, status', '角色表，role_code含STUDENT/TEACHER/LEADER/ADMIN'],
        ['identity_registry', 'id, identity_number, name, id_card_last6, role_id, department, major, grade, class_name, status', '身份注册库，由管理员预录入，status=0表示未注册、1表示已注册'],
    ]
)
doc.add_paragraph()

add_heading('4.1.3 接口设计', 2)
add_table(
    ['方法', 'URL', '请求体', '响应', '说明'],
    [
        ['POST', '/auth/login', '{identityNumber, password}', '{token, userId, name, roleCode, roleName}', '用户登录'],
        ['POST', '/auth/register', '{identityNumber, name, idCardLast6, password, phone, email}', '{code:200}', '用户注册（身份核验通过后创建账号）'],
        ['GET', '/user/profile', 'Header: Authorization Bearer xxx', '{16字段个人信息}', '查询个人信息（需JWT）'],
        ['PUT', '/user/profile', 'Header + 可修改字段', '{code:200}', '修改个人信息（需JWT）'],
    ]
)
doc.add_paragraph()

add_heading('4.2 模块间交互设计', 2)
add_para('模块间交互遵循"高内聚、低耦合"原则：')
add_para('• 通知公告查询：notice模块独立，不依赖其他模块，从JWT中获取角色信息进行公告定向推送')
add_para('• 课程-选课-成绩链路：course模块内部闭环——course(课程定义) → course_schedule(排课) → student_course(选课+成绩)，通过student_id关联user表')
add_para('• 考试-评教关联：exam_arrangement关联course_id获取课程信息，teaching_evaluation关联student_id/teacher_id/course_id')
add_para('• 一卡通-缴费：card模块下 card(卡片主表) + card_transaction(消费流水) + payment_record(缴费记录)，均通过user_id关联')
add_para('• 统计报表(领导端)：system模块聚合查询多个表（user + identity_registry + payment_record + teaching_evaluation + fixed_asset），通过Service层注入各模块Repository实现')
add_para('• AI助手：ai_knowledge_base预存知识条目（按角色+意图分类），用户提问时匹配关键词→返回答案模板+跳转链接；对话记录存入ai_conversation_history')

doc.add_page_break()

# ============ 5. 非功能需求设计 ============
add_heading('5. 非功能需求设计', 1)

add_heading('5.1 安全性设计', 2)
add_para('• 认证安全：JWT Token有效期24小时，Token含userId+roleCode，每次请求由JwtFilter校验签名和过期时间')
add_para('• 密码安全：BCrypt自适应哈希，salt随机生成，不可逆')
add_para('• 防SQL注入：全部使用JPA参数化查询（@Query + :param），无拼接SQL')
add_para('• 操作审计：login_log记录每次登录的IP、设备、时间、成功/失败状态')
add_para('• 注册安全：注册时必须通过identity_registry身份比对（学号+姓名+身份证后6位三方匹配）')
add_para('• 角色权限：每个API接口对应的Controller方法根据URL前缀区分角色（/admin/*仅管理员可访问），后续扩展到方法级注解鉴权')

add_heading('5.2 性能设计', 2)
add_para('• 数据库连接池：HikariCP，最大连接数20，最小空闲5，连接超时30秒')
add_para('• JPA懒加载：@OneToMany/@ManyToOne默认FetchType.LAZY，避免N+1查询')
add_para('• 前端打包优化：Vite Tree Shaking + 代码分割，首屏加载≤2秒')
add_para('• 静态资源：CSS/JS/图片经Vite处理后带hash，启用浏览器缓存')

add_heading('5.3 可靠性设计', 2)
add_para('• 全局异常处理：@RestControllerAdvice统一捕获RuntimeException，返回Result.fail(msg)')
add_para('• 事务管理：涉及多表操作的Service方法加@Transactional，异常时自动回滚')
add_para('• 前端容错：Axios响应拦截器统一处理401（跳登录页）、500（提示错误）、网络异常（提示重试）')
add_para('• 数据完整性：数据库外键约束 + JPA @Column(nullable=false)双重保障')

add_heading('5.4 可维护性设计', 2)
add_para('• 模块化包结构：controller/service/repository/entity/dto按模块分包，职责清晰')
add_para('• DTO集中管理：Request.java和Response.java内部分别定义所有接口的请求/响应DTO（内部类方式，减少文件数量）')
add_para('• 统一响应格式：Result<T>泛型类封装 {code, message, data}，全局统一')
add_para('• 代码注释：关键业务逻辑标注中文注释，说明业务含义和处理边界')

# ============ 6. 数据架构设计 ============
add_heading('6. 数据架构设计', 1)

add_heading('6.1 核心数据模型', 2)
add_para('系统共37张数据库表，分为10个业务域：')

add_table(
    ['业务域', '表名', '核心关系'],
    [
        ['用户权限(10张)', 'user, user_profile, user_role, role, role_permission, permission, identity_registry, register_application, login_log, password_history', 'user -< user_role >- role; user ||--|| user_profile; role -< role_permission >- permission'],
        ['通知公告(2张)', 'announcement, news', '独立表，publisher_id关联user.id'],
        ['课程成绩(5张)', 'course, course_schedule, student_course, student_retake, score_change_log', 'course -< course_schedule; course_schedule -< student_course; course -< student_retake'],
        ['考试评教(4张)', 'exam_arrangement, level_exam, level_exam_registration, teaching_evaluation', 'exam_arrangement关联course_id; level_exam -< level_exam_registration; teaching_evaluation关联student/teacher/course'],
        ['一卡通(3张)', 'card, card_transaction, payment_record', 'card ||--|| user; card -< card_transaction; payment_record关联student_id'],
        ['学生服务(4张)', 'student_application, competition_registration, consultation_appointment, practice_project_registration', '各表独立，通过student_id关联user'],
        ['教师发展(2张)', 'scientific_research, graduation_design', 'scientific_research关联teacher_id; graduation_design关联student_id + teacher_id'],
        ['公文计划(2张)', 'document, work_plan', '独立表，通过user_id关联'],
        ['系统管理(4张)', 'sys_config, fixed_asset, lab, login_log', '独立表'],
        ['AI助手(2张)', 'ai_knowledge_base, ai_conversation_history', 'ai_conversation_history关联user_id'],
    ]
)
doc.add_paragraph()

add_heading('6.2 关键关联关系', 2)
add_para('• 用户-角色：多对多（user_role中间表），一个用户可拥有多个角色')
add_para('• 角色-权限：多对多（role_permission中间表），支持RBAC灵活授权')
add_para('• 课程-学生：多对多（student_course中间表，通过course_schedule关联具体教学班）')
add_para('• 教师-课程：一对多（一位教师可教授多门课程）')
add_para('• 一卡通-用户：一对一（每位用户一张校园卡）')

add_heading('6.3 数据存储方案', 2)
add_para('• 主数据库：MySQL 8.4.5 InnoDB引擎，字符集utf8mb4')
add_para('• 密码字段：BCrypt哈希，长度255，不可逆')
add_para('• 大文本字段：公告内容、新闻内容、AI回答等使用LONGTEXT类型')
add_para('• 时间字段：统一使用datetime类型，Java侧对应LocalDateTime')
add_para('• 金额字段：使用decimal(10,2)或decimal(12,2)，避免浮点精度问题')

doc.add_page_break()

# ============ 7. 部署架构设计 ============
add_heading('7. 部署架构设计', 1)

add_heading('7.1 开发环境部署', 2)
add_para('开发阶段采用前后端独立部署：')
add_para('• 后端：IDE (IntelliJ IDEA) 直接运行 SmartcampusApplication.java → 监听 localhost:8080')
add_para('• 前端：终端执行 npm run dev → Vite Dev Server 监听 localhost:5173')
add_para('• 数据库：MySQL 8.4.5 Windows服务 MySQL8 → 监听 localhost:3306')
add_para('• 代理配置：Vite proxy 将 /api/* 请求转发至 localhost:8080，解决跨域问题')

add_heading('7.2 生产环境部署（规划）', 2)
add_para('• 前端：npm run build → 生成 dist/ 静态文件 → 部署至 Nginx')
add_para('• 后端：mvn package → target/*.jar → java -jar 运行')
add_para('• Nginx配置：80端口服务前端静态文件，/api/ 反向代理至后端8080端口')
add_para('• 数据库：独立MySQL服务器，定期备份')

add_heading('7.3 系统功能结构图（文字描述）', 2)
add_para('系统首页 = 登录页 → 根据角色跳转至对应的Dashboard首页：', bold=True)
add_para('')
add_para('学生端：Dashboard → [个人信息 | 公告通知 | 新闻浏览 | 我的课表 | 成绩查询 | 考试安排 | 网上评教 | 补考重修 | 等级考试报名 | 一卡通 | 缴费信息 | 学生服务(奖助贷/竞赛/咨询/实践) | AI助手]')
add_para('')
add_para('教师端：Dashboard → [个人信息 | 公告通知 | 任课信息 | 成绩录入 | 评教结果 | 科研管理 | 毕设管理 | 考试安排 | 待办公文 | 班级名单 | 奖助贷管理 | AI助手]')
add_para('')
add_para('领导端：Dashboard → [个人信息 | 公告通知 | 通讯录 | 工作计划 | 公文审批 | 科研总览 | 固定资产统计 | 新生报到统计 | 缴费统计 | 学籍异动 | 在校生统计 | 评教统计 | AI助手]')
add_para('')
add_para('管理员端：Dashboard → [系统概览 | 公告管理 | 新闻管理 | 用户管理 | 注册审核 | 角色权限 | 课程管理 | 任课分配 | 选课管理 | 考试安排 | 一卡通同步 | 缴费管理 | 日志审计 | 系统配置 | 固定资产 | 实验室管理 | AI助手]')

# ============ 8. 参考资料 ============
add_heading('8. 参考资料', 1)
add_para('• 《Spring Boot 4.1.0 Reference Documentation》')
add_para('• 《Vue 3.5 官方文档》')
add_para('• 《Element Plus 组件库文档》')
add_para('• 《MySQL 8.4 Reference Manual》')
add_para('• 《智慧校园服务平台 - 完整功能清单（四角色 + AI助手）》内部文档')
add_para('• 《软件需求规约》内部文档')
add_para('• 《数据库设计说明书》内部文档')

# ============ 保存 ============
output_path = r'D:\software\smart-campus\5.系统架构设计说明书-智慧校园.docx'
doc.save(output_path)
print(f'✅ 文档已生成: {output_path}')
